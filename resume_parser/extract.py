"""
Text Extraction Module

This module handles extracting plain text from PDF and DOCX files.
We use pdfplumber as the primary PDF parser because it handles tables and 
layouts better than PyPDF2, with PyMuPDF (fitz) as fallback for complex PDFs.
For DOCX files, we use python-docx for better formatting preservation.
"""

import logging
from pathlib import Path
from typing import Optional

# PDF parsing libraries - pdfplumber for primary, PyMuPDF for fallback
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logging.warning("pdfplumber not available")

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logging.warning("PyMuPDF not available")

# DOCX parsing libraries
try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logging.warning("python-docx not available")

try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False
    logging.warning("docx2txt not available")


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract plain text from a PDF file.
    
    Uses pdfplumber as primary method (better table/layout handling)
    with PyMuPDF as fallback for complex PDFs that pdfplumber can't handle.
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        str: Extracted plain text from the PDF
        
    Raises:
        FileNotFoundError: If the file doesn't exist
        ValueError: If neither PDF library is available
        Exception: For other PDF parsing errors
    """
    if not Path(file_path).exists():
        raise FileNotFoundError(f"PDF file not found: {file_path}")
    
    if not PDFPLUMBER_AVAILABLE and not PYMUPDF_AVAILABLE:
        raise ValueError("No PDF parsing library available. Install pdfplumber or PyMuPDF.")
    
    text = ""
    
    # Try pdfplumber first (better for resumes with tables/columns)
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if text.strip():  # Successfully extracted text
                return text
                
        except Exception as e:
            logging.warning(f"pdfplumber failed for {file_path}: {e}")
    
    # Fallback to PyMuPDF if pdfplumber failed or unavailable
    if PYMUPDF_AVAILABLE:
        try:
            pdf_document = fitz.open(file_path)
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                text += page.get_text() + "\n"
            pdf_document.close()
            
            if text.strip():
                return text
                
        except Exception as e:
            logging.error(f"PyMuPDF also failed for {file_path}: {e}")
            raise
    
    raise Exception(f"Failed to extract text from PDF: {file_path}")


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract plain text from a DOCX file.
    
    Uses python-docx as primary method (preserves formatting better)
    with docx2txt as simple fallback.
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        str: Extracted plain text from the DOCX
        
    Raises:
        FileNotFoundError: If the file doesn't exist
        ValueError: If neither DOCX library is available
        Exception: For other DOCX parsing errors
    """
    if not Path(file_path).exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")
    
    if not PYTHON_DOCX_AVAILABLE and not DOCX2TXT_AVAILABLE:
        raise ValueError("No DOCX parsing library available. Install python-docx or docx2txt.")
    
    # Try python-docx first (better formatting preservation)
    if PYTHON_DOCX_AVAILABLE:
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables (common in resumes)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            if text.strip():
                return text
                
        except Exception as e:
            logging.warning(f"python-docx failed for {file_path}: {e}")
    
    # Fallback to docx2txt (simpler extraction)
    if DOCX2TXT_AVAILABLE:
        try:
            text = docx2txt.process(file_path)
            if text and text.strip():
                return text
        except Exception as e:
            logging.error(f"docx2txt also failed for {file_path}: {e}")
            raise
    
    raise Exception(f"Failed to extract text from DOCX: {file_path}")


def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from a file based on its extension.
    
    Convenience function that automatically detects file type
    and calls the appropriate extraction function.
    
    Args:
        file_path (str): Path to the file (PDF or DOCX)
        
    Returns:
        str: Extracted plain text
        
    Raises:
        ValueError: For unsupported file types
    """
    file_path = Path(file_path)
    
    if file_path.suffix.lower() == '.pdf':
        return extract_text_from_pdf(str(file_path))
    elif file_path.suffix.lower() in ['.docx', '.doc']:
        return extract_text_from_docx(str(file_path))
    else:
        raise ValueError(f"Unsupported file type: {file_path.suffix}")
